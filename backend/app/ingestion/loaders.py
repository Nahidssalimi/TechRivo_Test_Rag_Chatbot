"""
Document loaders for various file formats
Handles PDF, CSV, TXT, DOCX files
"""

import os
from typing import List, Dict
import PyPDF2
import pandas as pd
from docx import Document as DocxDocument


class DocumentLoader:
    """Handles loading documents from various file formats"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.csv', '.txt', '.docx']
    
    def load_pdf(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and extract text from PDF files
        Returns list of dicts with page content and metadata
        """
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():  # Only add non-empty pages
                        documents.append({
                            'content': text,
                            'metadata': {
                                'source': file_path,
                                'page': page_num + 1,
                                'total_pages': total_pages,
                                'type': 'pdf'
                            }
                        })
        except Exception as e:
            print(f"Error loading PDF {file_path}: {str(e)}")
            
        return documents
    
    def load_csv(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and extract text from CSV files
        Converts each row to a text representation
        """
        documents = []
        
        try:
            df = pd.read_csv(file_path)
            
            # Convert entire dataframe to text representation
            full_text = df.to_string(index=False)
            documents.append({
                'content': full_text,
                'metadata': {
                    'source': file_path,
                    'rows': len(df),
                    'columns': list(df.columns),
                    'type': 'csv'
                }
            })
            
            # Also create individual documents for each row (optional, good for detailed search)
            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                documents.append({
                    'content': row_text,
                    'metadata': {
                        'source': file_path,
                        'row': idx + 1,
                        'type': 'csv_row'
                    }
                })
                
        except Exception as e:
            print(f"Error loading CSV {file_path}: {str(e)}")
            
        return documents
    
    def load_txt(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and extract text from TXT files
        """
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                if content.strip():
                    documents.append({
                        'content': content,
                        'metadata': {
                            'source': file_path,
                            'type': 'txt'
                        }
                    })
        except Exception as e:
            print(f"Error loading TXT {file_path}: {str(e)}")
            
        return documents
    
    def load_docx(self, file_path: str) -> List[Dict[str, str]]:
        """
        Load and extract text from DOCX files
        """
        documents = []
        
        try:
            doc = DocxDocument(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            content = "\n".join(full_text)
            
            if content.strip():
                documents.append({
                    'content': content,
                    'metadata': {
                        'source': file_path,
                        'paragraphs': len(doc.paragraphs),
                        'type': 'docx'
                    }
                })
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {str(e)}")
            
        return documents
    
    def load_document(self, file_path: str) -> List[Dict[str, str]]:
        """
        Auto-detect file type and load accordingly
        """
        _, ext = os.path.splitext(file_path.lower())
        
        if ext == '.pdf':
            return self.load_pdf(file_path)
        elif ext == '.csv':
            return self.load_csv(file_path)
        elif ext == '.txt':
            return self.load_txt(file_path)
        elif ext == '.docx':
            return self.load_docx(file_path)
        else:
            print(f"Unsupported file type: {ext}")
            return []
    
    def load_directory(self, directory_path: str) -> List[Dict[str, str]]:
        """
        Load all supported documents from a directory
        """
        all_documents = []
        
        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                _, ext = os.path.splitext(file.lower())
                
                if ext in self.supported_extensions:
                    print(f"Loading: {file_path}")
                    docs = self.load_document(file_path)
                    all_documents.extend(docs)
        
        return all_documents
